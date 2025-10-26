import os
import uuid
import time
from datetime import datetime, timedelta
from typing import List, Dict, Any, Optional
from fastapi import FastAPI, File, UploadFile, HTTPException, Form, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from pydantic import BaseModel
import PyPDF2
import docx
import csv
import io
from pymongo import MongoClient
from sentence_transformers import SentenceTransformer
import google.generativeai as genai
import numpy as np
import re
import hashlib
from bson import ObjectId
from contextlib import asynccontextmanager
import logging
import asyncio
from spellchecker import SpellChecker
import psutil
import pandas as pd
from collections import defaultdict
import matplotlib.pyplot as plt
import seaborn as sns
import json
import sys

# Load environment variables
from dotenv import load_dotenv
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize spell checker
spell = SpellChecker()

# Performance tracking variables
performance_metrics = {
    "query_history": [],
    "upload_history": [],
    "system_metrics": [],
    "start_time": datetime.utcnow()
}

class PerformanceMetrics:
    @staticmethod
    def track_query(query: str, response_time: float, confidence: float, 
                   relevant_docs: int, total_docs: int, success: bool):
        """Track query performance metrics"""
        metric = {
            "timestamp": datetime.utcnow().isoformat(),
            "query": query,
            "response_time_ms": response_time,
            "confidence": confidence,
            "relevant_docs": relevant_docs,
            "total_docs": total_docs,
            "success": success,
            "query_length": len(query)
        }
        performance_metrics["query_history"].append(metric)
        
        # Keep only last 1000 queries
        if len(performance_metrics["query_history"]) > 1000:
            performance_metrics["query_history"] = performance_metrics["query_history"][-1000:]

    @staticmethod
    def track_upload(filename: str, chunks_processed: int, chunks_failed: int, processing_time: float):
        """Track upload performance metrics"""
        metric = {
            "timestamp": datetime.utcnow().isoformat(),
            "filename": filename,
            "chunks_processed": chunks_processed,
            "chunks_failed": chunks_failed,
            "processing_time_sec": processing_time,
            "success_rate": chunks_processed / (chunks_processed + chunks_failed) if chunks_processed + chunks_failed > 0 else 0
        }
        performance_metrics["upload_history"].append(metric)

    @staticmethod
    def get_system_metrics():
        """Get current system performance metrics"""
        process = psutil.Process()
        memory_info = process.memory_info()
        
        return {
            "timestamp": datetime.utcnow().isoformat(),
            "memory_usage_mb": memory_info.rss / 1024 / 1024,
            "memory_percent": process.memory_percent(),
            "cpu_percent": process.cpu_percent(),
            "total_queries": len(performance_metrics["query_history"]),
            "total_uploads": len(performance_metrics["upload_history"]),
            "uptime_seconds": (datetime.utcnow() - performance_metrics["start_time"]).total_seconds()
        }

    @staticmethod
    def calculate_kpis(timeframe_hours: int = 24):
        """Calculate Key Performance Indicators"""
        cutoff_time = datetime.utcnow() - timedelta(hours=timeframe_hours)
        
        recent_queries = [
            q for q in performance_metrics["query_history"]
            if datetime.fromisoformat(q["timestamp"]) > cutoff_time
        ]
        
        if not recent_queries:
            return {
                "total_queries": 0,
                "avg_response_time": 0,
                "success_rate": 0,
                "avg_confidence": 0,
                "avg_relevant_docs": 0
            }
        
        successful_queries = [q for q in recent_queries if q["success"]]
        
        return {
            "total_queries": len(recent_queries),
            "avg_response_time": sum(q["response_time_ms"] for q in recent_queries) / len(recent_queries),
            "success_rate": len(successful_queries) / len(recent_queries),
            "avg_confidence": sum(q["confidence"] for q in successful_queries) / len(successful_queries) if successful_queries else 0,
            "avg_relevant_docs": sum(q["relevant_docs"] for q in successful_queries) / len(successful_queries) if successful_queries else 0
        }

# Cost tracking
class CostCalculator:
    # Approximate costs (adjust based on your actual usage)
    GEMINI_COST_PER_1K_TOKENS = 0.000125  # $0.000125 per 1K tokens for Gemini Flash
    EMBEDDING_COST_PER_1K = 0.0001        # $0.0001 per 1K embeddings
    
    @staticmethod
    def estimate_gemini_cost(text: str) -> float:
        """Estimate Gemini API cost based on text length"""
        # Rough estimate: 1 token ≈ 4 characters
        estimated_tokens = len(text) / 4
        cost = (estimated_tokens / 1000) * CostCalculator.GEMINI_COST_PER_1K_TOKENS
        return cost
    
    @staticmethod
    def estimate_embedding_cost(text: str) -> float:
        """Estimate embedding generation cost"""
        # Based on approximate costs for sentence transformers
        return CostCalculator.EMBEDDING_COST_PER_1K * (len(text) / 1000)
    
    @staticmethod
    def calculate_total_cost():
        """Calculate total estimated cost from query history"""
        total_cost = 0.0
        
        for query in performance_metrics["query_history"]:
            # Estimate Gemini cost for query + answer (approximate)
            query_text = query["query"]
            estimated_answer_length = 500  # Average answer length
            total_cost += CostCalculator.estimate_gemini_cost(query_text + " " * estimated_answer_length)
        
        return total_cost

# Initialize FastAPI app with lifespan management
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    await initialize_services()
    yield
    # Shutdown
    await close_connections()

app = FastAPI(
    title="Course Q&A API",
    description="Course Q&A Chatbot with MongoDB Semantic Search",
    version="4.8.0",
    lifespan=lifespan
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000", "http://127.0.0.1:3000", "http://localhost:3001"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Load environment variables from .env file
MONGO_URI = os.getenv("MONGO_URI")
MONGO_DB = os.getenv("MONGO_DB", "courseqa")
MONGO_COLLECTION = os.getenv("MONGO_COLLECTION", "documents")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")
PORT = int(os.getenv("PORT", "8000"))

# Global variables for services
embedding_model = None
mongo_client = None
documents_collection = None
services_initialized = False
gemini_model = None

async def initialize_services():
    """Initialize all services"""
    global services_initialized, gemini_model
    
    logger.info("🚀 Starting service initialization...")
    
    # Initialize MongoDB
    await initialize_mongodb()
    
    # Initialize ML models
    await initialize_ml_models()
    
    # Initialize Gemini model
    await initialize_gemini()
    
    services_initialized = True
    logger.info("✅ All services initialized successfully!")

async def initialize_mongodb(max_retries=3, retry_delay=2):
    """Initialize MongoDB connection with retry logic"""
    global mongo_client, documents_collection
    
    if not MONGO_URI:
        logger.error("❌ MongoDB URI not configured")
        mongo_client = None
        documents_collection = None
        return False
        
    for attempt in range(max_retries):
        try:
            logger.info(f"🔧 Attempting MongoDB connection (attempt {attempt + 1}/{max_retries})...")
            
            mongo_client = MongoClient(
                MONGO_URI, 
                serverSelectionTimeoutMS=15000,
                connectTimeoutMS=15000,
                socketTimeoutMS=15000,
                retryWrites=True,
                retryReads=True
            )
            
            # Test connection
            mongo_client.admin.command('ping')
            logger.info("✅ MongoDB connection test successful")
            
            # Get database and collection
            db = mongo_client[MONGO_DB]
            documents_collection = db[MONGO_COLLECTION]
            
            # Create indexes
            await create_mongodb_indexes()
            
            logger.info("✅ MongoDB initialized successfully")
            return True
            
        except Exception as e:
            logger.error(f"❌ MongoDB connection attempt {attempt + 1} failed: {e}")
            if attempt < max_retries - 1:
                logger.info(f"🔄 Retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
            else:
                logger.error("❌ All MongoDB connection attempts failed")
                mongo_client = None
                documents_collection = None
                return False

async def create_mongodb_indexes():
    """Create MongoDB indexes with error handling"""
    try:
        if documents_collection is None:
            return
            
        # Create text index
        try:
            documents_collection.create_index([("content", "text")])
            logger.info("✅ Created text index on 'content' field")
        except Exception as e:
            logger.info(f"✅ Text index already exists or couldn't be created: {e}")
        
        # Create chunk_id index
        try:
            documents_collection.create_index([("chunk_id", 1)], unique=True)
            logger.info("✅ Created unique index on 'chunk_id' field")
        except Exception as e:
            logger.info(f"✅ chunk_id index already exists or couldn't be created: {e}")
        
        # Create embedding index for semantic search
        try:
            documents_collection.create_index([("embedding", "2dsphere")])
            logger.info("✅ Created 2dsphere index on 'embedding' field for semantic search")
        except Exception as e:
            logger.info(f"✅ Embedding index already exists or couldn't be created: {e}")
            
    except Exception as e:
        logger.error(f"⚠️  Error creating indexes: {e}")

async def initialize_ml_models():
    """Initialize ML models"""
    global embedding_model
    
    try:
        logger.info("🔧 Loading embedding model...")
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        logger.info("✅ Embedding model loaded successfully")
        
        # Test the embedding model
        test_embedding = embedding_model.encode("test")
        logger.info(f"✅ Embedding model test successful - Output dimension: {len(test_embedding)}")
        
    except Exception as e:
        logger.error(f"❌ Embedding model loading error: {e}")
        embedding_model = None

async def initialize_gemini():
    """Initialize Google Gemini AI model"""
    global gemini_model
    
    if not GOOGLE_API_KEY:
        logger.warning("⚠️  Google AI Studio API key not configured")
        gemini_model = None
        return
        
    try:
        genai.configure(api_key=GOOGLE_API_KEY)
        gemini_model = genai.GenerativeModel('gemini-2.0-flash')
        logger.info("✅ Google Gemini model initialized")
        
        # Test the model
        try:
            response = gemini_model.generate_content("Hello")
            logger.info("✅ Gemini model test successful")
        except Exception as e:
            logger.warning(f"⚠️  Gemini model test failed: {e}")
        
    except Exception as e:
        logger.error(f"❌ Gemini initialization error: {e}")
        gemini_model = None

async def close_connections():
    """Close all connections"""
    if mongo_client:
        mongo_client.close()
        logger.info("✅ MongoDB connection closed")
    logger.info("✅ All connections closed")

def check_database_available():
    """Check if database is available"""
    if documents_collection is None:
        return False
    try:
        documents_collection.database.command('ping')
        return True
    except Exception as e:
        logger.error(f"Database connection test failed: {e}")
        return False

# Pydantic models
class AnswerRequest(BaseModel):
    query: str
    lang: str = "en"
    top_k: int = 5

class AnswerResponse(BaseModel):
    answer: str
    citations: List[Dict[str, Any]]
    confidence: float
    document_count: int
    relevant_docs: int
    latency_ms: float

class UploadResponse(BaseModel):
    message: str
    document_id: str
    processing_type: str
    chunks_processed: int
    chunks_failed: int

class DocumentCountResponse(BaseModel):
    document_count: int

class SourceResponse(BaseModel):
    text: str
    metadata: Dict[str, Any]
    source_id: str

# New Pydantic models for metrics
class KPIsResponse(BaseModel):
    total_queries: int
    avg_response_time: float
    success_rate: float
    avg_confidence: float
    avg_relevant_docs: float
    timeframe_hours: int

class SystemMetricsResponse(BaseModel):
    memory_usage_mb: float
    memory_percent: float
    cpu_percent: float
    total_queries: int
    total_uploads: int
    uptime_seconds: float

class CostMetricsResponse(BaseModel):
    total_estimated_cost_usd: float
    cost_per_query: float
    total_queries: int
    cost_breakdown: Dict[str, float]

# Enhanced text processing functions
def normalize_text(text: str) -> str:
    """Normalize text for better matching - handle plurals, verb forms, etc."""
    if not text:
        return ""
    
    # Convert to lowercase
    text = text.lower()
    
    # Handle common pluralizations and variations
    variations = {
        r'navigations?\b': 'navigation',
        r'browsers?\b': 'browser',
        r'drivers?\b': 'driver',
        r'automations?\b': 'automation',
        r'clicks?\b': 'click',
        r'opens?\b': 'open',
        r'tabs?\b': 'tab',
        r'windows?\b': 'window',
        r'buttons?\b': 'button',
        r'forms?\b': 'form',
        r'elements?\b': 'element',
        r'pages?\b': 'page',
        r'websites?\b': 'website',
        r'applications?\b': 'application',
        r'methods?\b': 'method',
        r'functions?\b': 'function',
        r'variables?\b': 'variable',
        r'classes?\b': 'class',
        r'objects?\b': 'object',
    }
    
    for pattern, replacement in variations.items():
        text = re.sub(pattern, replacement, text)
    
    return text

def expand_query_terms(query: str) -> list[str]:
    """Expand query with synonyms and related terms"""
    base_terms = query.lower().split()
    expanded_terms = set(base_terms)
    
    # Common synonyms and related terms
    synonyms = {
        "navigate": ["navigation", "browse", "go to", "visit", "open"],
        "browser": ["web browser", "chrome", "firefox", "safari", "edge"],
        "driver": ["webdriver", "selenium driver", "browser driver"],
        "open": ["launch", "start", "load", "initialize"],
        "tab": ["browser tab", "new tab", "window tab"],
        "click": ["press", "select", "tap", "choose"],
        "find": ["locate", "search", "identify", "detect"],
        "get": ["retrieve", "obtain", "fetch", "acquire"],
        "how": ["method", "way", "process", "procedure"],
        "what": ["explain", "describe", "define", "meaning"],
        "selenium": ["web automation", "browser automation", "web testing"],
        "webdriver": ["browser driver", "selenium driver", "automation driver"],
    }
    
    for term in base_terms:
        if term in synonyms:
            expanded_terms.update(synonyms[term])
    
    return list(expanded_terms)

def calculate_semantic_similarity(text1: str, text2: str) -> float:
    """Calculate semantic similarity between two texts using word overlap and pattern matching"""
    if not text1 or not text2:
        return 0.0
    
    # Normalize both texts
    norm1 = normalize_text(text1)
    norm2 = normalize_text(text2)
    
    # Split into words
    words1 = set(re.findall(r'\w+', norm1))
    words2 = set(re.findall(r'\w+', norm2))
    
    if not words1 or not words2:
        return 0.0
    
    # Calculate Jaccard similarity
    intersection = len(words1.intersection(words2))
    union = len(words1.union(words2))
    
    if union == 0:
        return 0.0
    
    jaccard_similarity = intersection / union
    
    # Bonus for exact phrase matches
    exact_match_boost = 0.0
    if text1.lower() in text2.lower() or text2.lower() in text1.lower():
        exact_match_boost = 0.3
    
    # Bonus for shared important terms
    important_terms = ["navigation", "browser", "driver", "selenium", "webdriver", "click", "open", "tab", "window", "automation"]
    shared_important = len([term for term in important_terms if term in norm1 and term in norm2])
    important_boost = shared_important * 0.1
    
    return min(jaccard_similarity + exact_match_boost + important_boost, 1.0)

# Enhanced Question Analysis with improved typo handling
def analyze_question_type(query: str) -> Dict[str, Any]:
    """Analyze the question to determine its type and requirements with enhanced typo handling"""
    query_lower = query.lower().strip()
    
    # Enhanced typo handling for common misspellings
    typo_corrections = {
        "id": "is", "hwodo": "how do", "whos": "who is", "wats": "what is",
        "navigations": "navigation", "selenimum": "selenium", "webdriverr": "webdriver",
        "brower": "browser", "drivr": "driver", "automat": "automate", "expln": "explain",
        "hw": "how", "wat": "what", "wen": "when", "wer": "where", "y": "why"
    }
    
    for typo, correction in typo_corrections.items():
        query_lower = query_lower.replace(typo, correction)
    
    # Auto-correct spelling using spell checker
    words = query_lower.split()
    corrected_words = []
    for word in words:
        if len(word) > 2:  # Only correct words longer than 2 characters
            correction = spell.correction(word)
            if correction and correction != word:
                corrected_words.append(correction)
                logger.info(f"🔧 Corrected '{word}' to '{correction}'")
            else:
                corrected_words.append(word)
        else:
            corrected_words.append(word)
    
    query_lower = " ".join(corrected_words)
    
    question_types = {
        "who": any(keyword in query_lower for keyword in ["who is", "who are", "who was", "who's", "this person"]),
        "what": any(keyword in query_lower for keyword in ["what is", "what are", "what does", "what's", "explain"]),
        "how": any(keyword in query_lower for keyword in ["how to", "how do", "how does", "how can", "how i"]),
        "why": any(keyword in query_lower for keyword in ["why", "reason", "purpose"]),
        "contact": any(keyword in query_lower for keyword in ["contact", "email", "phone", "mobile", "number", "reach", "get in touch"]),
        "skills": any(keyword in query_lower for keyword in ["skills", "technologies", "languages", "frameworks", "proficient in"]),
        "projects": any(keyword in query_lower for keyword in ["projects", "project", "work on", "built", "developed", "explain project"]),
        "education": any(keyword in query_lower for keyword in ["education", "degree", "college", "university", "study"]),
        "technical": any(keyword in query_lower for keyword in ["code", "program", "script", "function", "method", "class", "variable"]),
        "navigation": any(keyword in query_lower for keyword in ["navigate", "navigation", "browser", "url", "website", "webpage"]),
    }
    
    # Remove name detection - search names in documents instead
    name_detected = False
    
    # Determine primary question type
    primary_type = next((q_type for q_type, matches in question_types.items() if matches), "general")
    
    # Override for specific queries
    if "mobile number" in query_lower or "phone number" in query_lower:
        primary_type = "contact"
    
    return {
        "type": primary_type,
        "requires_person_info": question_types["who"] or name_detected,
        "requires_contact": question_types["contact"],
        "requires_skills": question_types["skills"],
        "requires_projects": question_types["projects"],
        "requires_education": question_types["education"],
        "requires_technical": question_types["technical"],
        "requires_navigation": question_types["navigation"],
        "name_detected": name_detected,
        "original_query": query,
        "processed_query": query_lower
    }

# IMPROVED Content Extraction Functions
def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file with improved parsing"""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        
        for page_num, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            if page_text:
                # Clean up the text
                page_text = re.sub(r'\s+', ' ', page_text).strip()
                text += f"Page {page_num + 1}: {page_text}\n\n"
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No readable text found in PDF")
            
        return text.strip()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading PDF: {str(e)}")

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file - IMPROVED VERSION"""
    try:
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        text = ""
        
        # Extract paragraphs with better formatting
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                # Preserve some formatting
                para_text = paragraph.text.strip()
                if paragraph.style.name.startswith('Heading'):
                    text += f"\n# {para_text}\n\n"
                else:
                    text += para_text + "\n"
        
        # Extract tables with better structure
        for table_idx, table in enumerate(doc.tables):
            text += f"\n[Table {table_idx + 1}]\n"
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    text += f"Row {row_idx + 1}: {' | '.join(row_data)}\n"
            text += "\n"
        
        # Extract headers and footers
        try:
            for section in doc.sections:
                # Header
                header = section.header
                if header:
                    for paragraph in header.paragraphs:
                        if paragraph.text and paragraph.text.strip():
                            text += f"[Header] {paragraph.text.strip()}\n"
                # Footer
                footer = section.footer
                if footer:
                    for paragraph in footer.paragraphs:
                        if paragraph.text and paragraph.text.strip():
                            text += f"[Footer] {paragraph.text.strip()}\n"
        except Exception as e:
            logger.warning(f"Could not extract headers/footers: {e}")
        
        # Clean up the text
        text = re.sub(r'\n{3,}', '\n\n', text)  # Replace multiple newlines
        text = text.strip()
        
        if not text:
            raise HTTPException(status_code=400, detail="No readable text found in DOCX document")
            
        logger.info(f"✅ Extracted {len(text)} characters from DOCX document")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting from DOCX: {e}")
        raise HTTPException(status_code=400, detail=f"Error reading DOCX document: {str(e)}")

def extract_text_from_xlsx(file_content: bytes) -> str:
    """Extract text from Excel files (XLSX)"""
    try:
        excel_file = io.BytesIO(file_content)
        
        # Read all sheets
        excel_data = pd.read_excel(excel_file, sheet_name=None)
        text = ""
        
        for sheet_name, df in excel_data.items():
            text += f"\n[Sheet: {sheet_name}]\n"
            
            # Add column headers
            if not df.empty:
                columns = df.columns.tolist()
                text += f"Columns: {', '.join(map(str, columns))}\n\n"
                
                # Add sample data (first 20 rows to avoid too much text)
                for idx, row in df.head(20).iterrows():
                    row_text = []
                    for col in df.columns:
                        cell_value = row[col]
                        if pd.notna(cell_value) and str(cell_value).strip():
                            row_text.append(f"{col}: {str(cell_value).strip()}")
                    if row_text:
                        text += f"Row {idx + 1}: {' | '.join(row_text)}\n"
                
                # Add summary
                text += f"\nTotal rows in '{sheet_name}': {len(df)}\n"
                text += "-" * 50 + "\n"
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No readable data found in Excel file")
            
        logger.info(f"✅ Extracted data from Excel file with {len(excel_data)} sheets")
        return text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting from Excel: {e}")
        raise HTTPException(status_code=400, detail=f"Error reading Excel file: {str(e)}")

def extract_text_from_csv(file_content: bytes) -> str:
    """Extract text from CSV file with better formatting"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']
        
        for encoding in encodings:
            try:
                csv_content = file_content.decode(encoding)
                csv_file = io.StringIO(csv_content)
                
                # Try to detect delimiter
                sample = csv_content[:1000]
                sniffer = csv.Sniffer()
                delimiter = sniffer.sniff(sample).delimiter
                
                csv_reader = csv.reader(csv_file, delimiter=delimiter)
                text = ""
                headers = []
                
                for i, row in enumerate(csv_reader):
                    if row and any(cell.strip() for cell in row):
                        clean_row = [cell.strip() for cell in row if cell.strip()]
                        if i == 0:  # Header row
                            headers = clean_row
                            text += f"Headers: {', '.join(headers)}\n\n"
                        else:
                            if headers:
                                # Create key-value pairs with headers
                                row_data = []
                                for j, value in enumerate(clean_row):
                                    if j < len(headers):
                                        row_data.append(f"{headers[j]}: {value}")
                                    else:
                                        row_data.append(value)
                                text += f"Row {i}: {' | '.join(row_data)}\n"
                            else:
                                text += f"Row {i}: {', '.join(clean_row)}\n"
                
                if not text.strip():
                    continue  # Try next encoding
                    
                return text.strip()
                
            except (UnicodeDecodeError, csv.Error, Exception):
                continue
        
        raise HTTPException(status_code=400, detail="Could not read CSV file with any supported encoding")
        
    except Exception as e:
        logger.error(f"Error reading CSV: {e}")
        raise HTTPException(status_code=400, detail=f"Error reading CSV file: {str(e)}")

def extract_text_from_txt(file_content: bytes) -> str:
    """Extract text from TXT file with encoding detection"""
    try:
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1', 'utf-16']
        
        for encoding in encodings:
            try:
                text = file_content.decode(encoding)
                # Clean up the text
                text = re.sub(r'\r\n', '\n', text)
                text = re.sub(r'\n\s*\n', '\n\n', text)
                if text.strip():
                    return text.strip()
            except UnicodeDecodeError:
                continue
        
        raise UnicodeDecodeError("Could not decode file with any supported encoding")
        
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Error reading text file: {str(e)}")

def intelligent_chunking(text: str, filename: str, chunk_size: int = 800, overlap: int = 100) -> List[Dict[str, Any]]:
    """Improved chunking that preserves complete sentences and paragraphs"""
    if not text or not text.strip():
        return []
    
    chunks = []
    
    # Normalize text
    text = re.sub(r'\r\n', '\n', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Split by major sections first (preserve document structure)
    section_pattern = r'(\n#{1,6}\s+.+|\n\s*(?:CHAPTER|SECTION|TOPIC|UNIT)\s+[\dIVX]+.*?|\n\s*\d+\.\s*[A-Z][^\n]+|\n\s*[A-Z][A-Z\s]{10,}[A-Z]\s*\n)'
    sections = re.split(section_pattern, text, flags=re.IGNORECASE)
    
    current_section = "Document Content"
    chunk_id = 0
    
    if len(sections) > 1:
        # Process with sections
        for i in range(0, len(sections), 2):
            if i < len(sections):
                if i + 1 < len(sections):
                    section_title = sections[i].strip() if sections[i].strip() else current_section
                    section_content = sections[i + 1]
                    
                    if section_title and len(section_title) > 2:
                        current_section = section_title
                    
                    if section_content.strip():
                        content_chunks = split_into_meaningful_chunks(section_content, chunk_size, overlap)
                        for chunk_content in content_chunks:
                            if chunk_content.strip():
                                chunk = create_chunk(chunk_content, filename, current_section, chunk_id)
                                chunks.append(chunk)
                                chunk_id += 1
    else:
        # No sections found, split the entire text
        content_chunks = split_into_meaningful_chunks(text, chunk_size, overlap)
        for chunk_content in content_chunks:
            if chunk_content.strip():
                chunk = create_chunk(chunk_content, filename, "Content", chunk_id)
                chunks.append(chunk)
                chunk_id += 1
    
    logger.info(f"Created {len(chunks)} chunks with average length {sum(len(c['content']) for c in chunks) // len(chunks) if chunks else 0} characters")
    return chunks

def split_into_meaningful_chunks(text: str, chunk_size: int, overlap: int) -> List[str]:
    """Split text into chunks while preserving sentence and paragraph boundaries"""
    chunks = []
    
    # First, split by paragraphs
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    
    current_chunk = ""
    
    for paragraph in paragraphs:
        # If current chunk + paragraph would exceed size and current chunk is not empty
        if len(current_chunk) + len(paragraph) + 2 > chunk_size and current_chunk:
            # Add the current chunk to chunks
            chunks.append(current_chunk.strip())
            
            # Start new chunk with overlap from previous chunk
            if overlap > 0:
                # Use last few sentences as overlap
                sentences = re.split(r'(?<=[.!?])\s+', current_chunk)
                overlap_text = ' '.join(sentences[-min(3, len(sentences)):])
                current_chunk = overlap_text + "\n\n" + paragraph if overlap_text else paragraph
            else:
                current_chunk = paragraph
        else:
            # Add paragraph to current chunk
            if current_chunk:
                current_chunk += "\n\n" + paragraph
            else:
                current_chunk = paragraph
    
    # Add the final chunk
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    # If chunks are still too large, split by sentences within paragraphs
    final_chunks = []
    for chunk in chunks:
        if len(chunk) > chunk_size * 1.5:
            # Split this chunk by sentences
            sentences = re.split(r'(?<=[.!?])\s+', chunk)
            temp_chunk = ""
            for sentence in sentences:
                if len(temp_chunk) + len(sentence) + 1 > chunk_size and temp_chunk:
                    final_chunks.append(temp_chunk.strip())
                    temp_chunk = sentence
                else:
                    if temp_chunk:
                        temp_chunk += " " + sentence
                    else:
                        temp_chunk = sentence
            if temp_chunk.strip():
                final_chunks.append(temp_chunk.strip())
        else:
            final_chunks.append(chunk)
    
    return final_chunks

def create_chunk(content: str, filename: str, section: str, chunk_id: int) -> Dict[str, Any]:
    """Create a standardized chunk object"""
    # Generate a unique chunk ID
    chunk_hash = hashlib.md5(f"{filename}_{section}_{chunk_id}_{content[:50]}".encode()).hexdigest()[:12]
    chunk_identifier = f"{filename}_{chunk_hash}"
    
    return {
        "chunk_id": chunk_identifier,
        "content": content,
        "section": section,
        "filename": filename,
        "page_number": 1,
        "created_at": datetime.utcnow(),
        "metadata": {
            "section": section,
            "filename": filename,
            "chunk_index": chunk_id,
            "content_length": len(content),
            "word_count": len(content.split()),
            "chunk_hash": chunk_hash
        }
    }

def generate_embedding(text: str) -> List[float]:
    """Generate embedding for text"""
    if embedding_model is None:
        logger.warning("⚠️  Embedding model not available, returning zero vector")
        return [0.0] * 384
    
    try:
        # Limit text length for embedding to avoid token limits
        if len(text) > 2000:
            text = text[:2000]
        embedding = embedding_model.encode(text).tolist()
        logger.info(f"✅ Generated embedding with {len(embedding)} dimensions")
        return embedding
    except Exception as e:
        logger.error(f"❌ Error generating embedding: {e}")
        return [0.0] * 384

async def gemini_chat_completion(prompt: str, max_tokens: int = 2000) -> Dict[str, Any]:
    """Google Gemini API integration with better error handling"""
    
    if not GOOGLE_API_KEY:
        return {"error": "Google AI Studio API key not configured"}
    
    if gemini_model is None:
        return {"error": "Gemini model not initialized"}
    
    try:
        # Generate content using Gemini API
        response = await asyncio.get_event_loop().run_in_executor(
            None,
            lambda: gemini_model.generate_content(
                prompt,
                generation_config=genai.types.GenerationConfig(
                    max_output_tokens=max_tokens,
                    temperature=0.1,  # Lower temperature for more factual answers
                    top_p=0.8,
                )
            )
        )
        
        if response.text:
            return {
                "choices": [
                    {
                        "message": {
                            "content": response.text
                        }
                    }
                ]
            }
        else:
            logger.warning("Empty response from Gemini API")
            return {"error": "Empty response from Gemini API"}
            
    except Exception as e:
        logger.error(f"Gemini API error: {e}")
        return {"error": f"Google Gemini API request failed: {str(e)}"}

def semantic_search(query: str, top_k: int = 20) -> List[Dict[str, Any]]:
    """Enhanced semantic search using MongoDB with vector similarity"""
    try:
        if documents_collection is None:
            logger.warning("❌ Documents collection is None")
            return []
        
        if embedding_model is None:
            logger.warning("❌ Embedding model not available, using text search")
            return text_search(query, top_k)
        
        # Generate query embedding
        query_embedding = generate_embedding(query)
        
        # Use cosine similarity search as fallback
        results = cosine_similarity_search(query_embedding, top_k)
        
        # Calculate enhanced relevance scores
        for result in results:
            content = result["content"]
            vector_score = result.get("score", 0.5)
            semantic_score = calculate_semantic_similarity(query, content)
            
            # Combine vector and semantic scores
            result["combined_score"] = (vector_score * 0.7 + semantic_score * 0.3)
            result["search_type"] = "vector_semantic"
        
        # Sort by combined score
        results.sort(key=lambda x: x.get("combined_score", 0), reverse=True)
        
        logger.info(f"📚 Semantic search found {len(results)} relevant chunks")
        return results[:top_k]
            
    except Exception as e:
        logger.error(f"Error in semantic search: {e}")
        return text_search(query, top_k)

def cosine_similarity_search(query_embedding: List[float], top_k: int) -> List[Dict[str, Any]]:
    """Fallback search using cosine similarity"""
    try:
        all_docs = list(documents_collection.find(
            {"embedding": {"$exists": True}}, 
            {"chunk_id": 1, "content": 1, "section": 1, "filename": 1, "page_number": 1, "metadata": 1, "embedding": 1}
        ).limit(1000))  # Limit for performance
        
        scored_docs = []
        for doc in all_docs:
            if "embedding" in doc and doc["embedding"]:
                doc_embedding = doc["embedding"]
                # Calculate cosine similarity
                similarity = np.dot(query_embedding, doc_embedding) / (
                    np.linalg.norm(query_embedding) * np.linalg.norm(doc_embedding)
                )
                doc["score"] = float(similarity)
                scored_docs.append(doc)
        
        scored_docs.sort(key=lambda x: x.get("score", 0), reverse=True)
        return scored_docs[:top_k]
        
    except Exception as e:
        logger.error(f"Cosine similarity search failed: {e}")
        return []

def text_search(query: str, top_k: int) -> List[Dict[str, Any]]:
    """Enhanced text-based search with better matching"""
    try:
        # Normalize and expand query
        normalized_query = normalize_text(query)
        expanded_terms = expand_query_terms(query)
        
        logger.info(f"🔍 Text search for: '{query}' -> Normalized: '{normalized_query}'")
        logger.info(f"📝 Expanded terms: {expanded_terms}")
        
        # Build search queries with multiple strategies
        search_queries = []
        
        # 1. Exact phrase match (boosted)
        if len(normalized_query.split()) > 1:
            exact_phrase = normalized_query
            search_queries.append({
                "filter": {"content": {"$regex": re.escape(exact_phrase), "$options": "i"}},
                "boost": 15.0,
                "type": "exact_phrase"
            })
        
        # 2. All expanded terms (AND)
        if expanded_terms:
            and_conditions = [{"content": {"$regex": re.escape(term), "$options": "i"}} for term in expanded_terms[:5]]
            if and_conditions:
                search_queries.append({
                    "filter": {"$and": and_conditions},
                    "boost": 8.0,
                    "type": "all_terms"
                })
        
        # 3. Any expanded terms (OR)
        if expanded_terms:
            or_conditions = [{"content": {"$regex": re.escape(term), "$options": "i"}} for term in expanded_terms[:8]]
            if or_conditions:
                search_queries.append({
                    "filter": {"$or": or_conditions},
                    "boost": 3.0,
                    "type": "any_terms"
                })
        
        # 4. Individual term matches
        individual_terms = normalized_query.split()
        for term in individual_terms[:4]:
            if len(term) > 2:  # Only search for terms longer than 2 characters
                search_queries.append({
                    "filter": {"content": {"$regex": re.escape(term), "$options": "i"}},
                    "boost": 2.0,
                    "type": f"term_{term}"
                })
        
        # Execute all search strategies
        all_results = []
        seen_chunks = set()
        
        for search_query in search_queries:
            try:
                results = list(documents_collection.find(
                    search_query["filter"],
                    {
                        "chunk_id": 1,
                        "content": 1,
                        "section": 1,
                        "filename": 1,
                        "page_number": 1,
                        "metadata": 1
                    }
                ).limit(top_k * 2))
                
                for result in results:
                    chunk_id = result["chunk_id"]
                    if chunk_id not in seen_chunks:
                        seen_chunks.add(chunk_id)
                        # Calculate enhanced relevance score
                        score = calculate_enhanced_relevance(result["content"], query, search_query["boost"])
                        result["score"] = score
                        result["search_type"] = search_query["type"]
                        all_results.append(result)
                        
            except Exception as e:
                logger.warning(f"Search query {search_query['type']} failed: {e}")
                continue
        
        # If no results found, try MongoDB text search
        if not all_results:
            logger.info("🔄 No results found, trying MongoDB text search...")
            try:
                text_results = list(documents_collection.find(
                    {"$text": {"$search": query}},
                    {
                        "chunk_id": 1,
                        "content": 1,
                        "section": 1,
                        "filename": 1,
                        "page_number": 1,
                        "metadata": 1,
                        "score": {"$meta": "textScore"}
                    }
                ).sort([("score", {"$meta": "textScore"})]).limit(top_k))
                
                for result in text_results:
                    chunk_id = result["chunk_id"]
                    if chunk_id not in seen_chunks:
                        seen_chunks.add(chunk_id)
                        result["score"] = result.get("score", 0.5) * 0.7  # Reduce score for text search
                        result["search_type"] = "text_search"
                        all_results.append(result)
            except Exception as e:
                logger.warning(f"MongoDB text search failed: {e}")
        
        # Sort by score and return top results
        all_results.sort(key=lambda x: x.get("score", 0), reverse=True)
        
        logger.info(f"📚 Text search found {len(all_results)} total results")
        return all_results[:top_k]
            
    except Exception as e:
        logger.error(f"Error in text search: {e}")
        return []

def calculate_enhanced_relevance(content: str, query: str, base_boost: float = 1.0) -> float:
    """Calculate enhanced relevance between content and query"""
    score = 0.0
    query_lower = query.lower()
    content_lower = content.lower()
    
    # Calculate semantic similarity
    semantic_similarity = calculate_semantic_similarity(query, content)
    score += semantic_similarity * 10.0 * base_boost
    
    # Exact phrase match (highest priority)
    if query_lower in content_lower:
        score += 20.0 * base_boost
    
    # Word overlap
    query_words = set(re.findall(r'\w+', query_lower))
    content_words = re.findall(r'\w+', content_lower)
    content_word_set = set(content_words)
    
    # Calculate word overlap
    matching_words = query_words.intersection(content_word_set)
    if matching_words:
        score += len(matching_words) * 3.0 * base_boost
    
    # Term frequency bonus
    for word in matching_words:
        term_freq = content_words.count(word)
        score += min(term_freq * 0.5, 5.0) * base_boost
    
    return min(score, 50.0)  # Cap score

def question_aware_retrieval(query: str, question_type: Dict[str, Any], top_k: int = 15) -> List[Dict[str, Any]]:
    """Enhanced retrieval that considers question type and handles variations"""
    
    # Use processed query for retrieval if available
    retrieval_query = question_type.get('processed_query', query)
    
    # Base semantic search
    results = semantic_search(retrieval_query, top_k * 2)
    
    # If we have few results, try with the original query as fallback
    if len(results) < 3 and question_type.get('processed_query') != query:
        logger.info("🔄 Few results found, trying with original query...")
        additional_results = semantic_search(question_type['original_query'], top_k)
        # Merge results
        seen_chunks = set(r["chunk_id"] for r in results)
        for result in additional_results:
            if result["chunk_id"] not in seen_chunks:
                results.append(result)
                seen_chunks.add(result["chunk_id"])
    
    # If still no results, try broader text search
    if not results:
        logger.info("🔄 No semantic results, trying broader text search...")
        results = text_search(query, top_k * 2)
    
    if not results:
        return []
    
    # Question-type specific boosting
    boosted_results = []
    for result in results:
        content = result["content"].lower()
        score = result.get('combined_score', result.get('score', 0.5))
        
        # Enhanced content analysis
        normalized_content = normalize_text(content)
        
        # Boost based on question type
        if question_type["requires_person_info"]:
            if any(keyword in normalized_content for keyword in ["student", "name", "person", "individual", "education", "contact"]):
                score *= 2.0
                
        elif question_type["requires_contact"]:
            if any(keyword in normalized_content for keyword in ["email", "phone", "mobile", "contact", "@", "gmail"]):
                score *= 3.0
                
        elif question_type["requires_skills"]:
            if any(keyword in normalized_content for keyword in ["skills", "technologies", "languages", "frameworks", "tools"]):
                score *= 2.5
                
        elif question_type["requires_projects"]:
            if any(keyword in normalized_content for keyword in ["projects", "project", "built", "developed", "created"]):
                score *= 3.0
                
        elif question_type["requires_education"]:
            if any(keyword in normalized_content for keyword in ["education", "degree", "college", "university", "institute"]):
                score *= 2.5
                
        elif question_type["requires_technical"]:
            if any(keyword in normalized_content for keyword in ["code", "program", "script", "function", "method", "class", "variable"]):
                score *= 2.0
                
        elif question_type["requires_navigation"]:
            if any(keyword in normalized_content for keyword in ["navigate", "navigation", "browser", "url", "website", "webpage"]):
                score *= 3.0
                
        result["question_aware_score"] = score
        boosted_results.append(result)
    
    # Sort by boosted score
    boosted_results.sort(key=lambda x: x.get("question_aware_score", 0), reverse=True)
    
    logger.info(f"🎯 Question-aware boosting applied, returning {len(boosted_results[:top_k])} results")
    return boosted_results[:top_k]

# Answer Generation Functions
async def generate_comprehensive_answer(query: str, context_chunks: List[Dict[str, Any]], question_type: Dict[str, Any]) -> Dict[str, Any]:
    """Generate comprehensive answer using Gemini API with retrieved context"""
    
    if not context_chunks:
        logger.warning("❌ No context chunks found for query")
        return await generate_fallback_answer(query)
    
    try:
        # Prepare context for Gemini
        context_text = ""
        citation_map = {}
        
        for i, chunk in enumerate(context_chunks[:10]):
            source_id = f"S{i+1}"
            context_text += f"[{source_id}]: {chunk['content']}\n\n"
            citation_map[source_id] = {
                "chunk_id": chunk["chunk_id"],
                "section": chunk.get("section", "Unknown"),
                "page_number": chunk.get("page_number", 1),
                "score": chunk.get('combined_score', chunk.get('score', 0.7))
            }
        
        # Create enhanced Gemini prompt
        system_prompt = f"""You are an expert educational assistant. Answer the user's question using ONLY the provided context from course documents.

CONTEXT DOCUMENTS:
{context_text}

USER QUESTION: {query}

CRITICAL INSTRUCTIONS:
1. ANSWER STRICTLY AND ONLY USING THE INFORMATION PROVIDED IN THE CONTEXT ABOVE
2. If the answer cannot be found in the context, say "I couldn't find specific information about this in the course materials."
3. Be specific, accurate, and detailed in your answer - provide complete explanations
4. Include relevant details, examples, or explanations from the context
5. If mentioning specific information, reference the source using the format [S1], [S2], etc.
6. Do not make up any information not present in the context
7. If the context contains multiple relevant points, include all of them
8. Structure your answer to be clear and helpful
9. Provide complete code examples or explanations when they appear in the context

ANSWER:"""
        
        # Use Gemini to generate answer
        response = await gemini_chat_completion(system_prompt, max_tokens=2500)
        
        if "error" in response:
            logger.warning(f"Gemini API failed, using enhanced fallback: {response['error']}")
            return generate_enhanced_structured_answer(query, context_chunks, question_type)
        
        answer_text = response["choices"][0]["message"]["content"]
        
        # Extract citations from answer text
        citations = extract_citations_from_answer(answer_text, citation_map)
        
        # If no citations found but we have context, add the top chunks as citations
        if not citations and context_chunks:
            citations = [{
                "source_id": chunk["chunk_id"],
                "span": f"pg{chunk.get('page_number', 1)}",
                "confidence": chunk.get('combined_score', chunk.get('score', 0.7)),
                "section": chunk.get("section", "Unknown"),
                "page_number": chunk.get("page_number", 1)
            } for chunk in context_chunks[:3]]
        
        confidence = calculate_enhanced_confidence(context_chunks, query, answer_text)
        
        return {
            "answer": answer_text,
            "citations": citations,
            "confidence": confidence
        }
        
    except Exception as e:
        logger.error(f"Error in Gemini answer generation: {e}")
        return generate_enhanced_structured_answer(query, context_chunks, question_type)

def extract_citations_from_answer(answer: str, citation_map: Dict[str, Any]) -> List[Dict[str, Any]]:
    """Extract citations from answer text based on [S1], [S2] markers"""
    citations = []
    
    # Find all citation markers in the answer
    citation_pattern = r'\[S(\d+)\]'
    matches = re.findall(citation_pattern, answer)
    
    for match in matches:
        source_id = f"S{match}"
        if source_id in citation_map:
            chunk_info = citation_map[source_id]
            citations.append({
                "source_id": chunk_info["chunk_id"],
                "span": f"pg{chunk_info['page_number']}",
                "confidence": chunk_info["score"],
                "section": chunk_info["section"],
                "page_number": chunk_info["page_number"]
            })
    
    # Remove duplicates
    unique_citations = []
    seen_ids = set()
    for citation in citations:
        if citation["source_id"] not in seen_ids:
            unique_citations.append(citation)
            seen_ids.add(citation["source_id"])
    
    return unique_citations

async def generate_fallback_answer(query: str) -> Dict[str, Any]:
    """Generate helpful fallback answer when no specific context is found"""
    # Try a broader search with relaxed criteria
    try:
        broader_results = text_search(query, top_k=25)
        if broader_results:
            # Use the broader results to generate an answer
            return generate_enhanced_structured_answer(query, broader_results[:8], {"type": "general"})
    except Exception as e:
        logger.warning(f"Broader search failed: {e}")
    
    return {
        "answer": "I couldn't find specific information about this in the uploaded course documents. The documents may not contain information about this specific topic. Please try asking about content that exists in your uploaded materials.",
        "citations": [],
        "confidence": 0.1
    }

def generate_enhanced_structured_answer(query: str, context_chunks: List[Dict[str, Any]], question_type: Dict[str, Any]) -> Dict[str, Any]:
    """Generate structured answer by analyzing content directly"""
    
    if not context_chunks:
        return {
            "answer": "I couldn't find specific information about this topic in the course materials.",
            "citations": [],
            "confidence": 0.1
        }
    
    # Extract key information from chunks
    relevant_content = []
    for chunk in context_chunks[:6]:
        content = chunk["content"]
        if calculate_semantic_similarity(query, content) > 0.3:
            relevant_content.append(content)
    
    if relevant_content:
        # Combine the most relevant parts
        combined_content = "\n\n".join(relevant_content[:4])
        
        # Try to extract a coherent answer
        sentences = re.split(r'[.!?]+', combined_content)
        relevant_sentences = []
        
        for sentence in sentences:
            if len(sentence.strip()) > 20:
                relevant_sentences.append(sentence.strip())
        
        if relevant_sentences:
            answer = "Based on the course materials:\n\n" + ". ".join(relevant_sentences[:8]) + "."
            confidence = 0.7
        else:
            # Use the beginning of the most relevant chunks
            previews = [chunk["content"][:300] + "..." for chunk in context_chunks[:3]]
            answer = "Here's relevant information from the course materials:\n\n" + "\n\n".join(previews)
            confidence = 0.5
    else:
        answer = "I found some related information in the documents but couldn't extract a specific answer to your question."
        confidence = 0.3
    
    # Generate citations
    citations = [{
        "source_id": chunk["chunk_id"],
        "span": f"pg{chunk.get('page_number', 1)}",
        "confidence": chunk.get('combined_score', chunk.get('score', 0.6)),
        "section": chunk.get("section", "Unknown"),
        "page_number": chunk.get("page_number", 1)
    } for chunk in context_chunks[:3]]
    
    return {
        "answer": answer,
        "citations": citations,
        "confidence": confidence
    }

def calculate_enhanced_confidence(context_chunks: List[Dict[str, Any]], query: str, answer: str) -> float:
    """Calculate confidence score based on multiple factors"""
    if not context_chunks:
        return 0.1
    
    # Base confidence from retrieval scores
    retrieval_scores = [chunk.get('combined_score', chunk.get('score', 0.3)) for chunk in context_chunks]
    avg_retrieval_score = sum(retrieval_scores) / len(retrieval_scores) if retrieval_scores else 0.3
    
    # Answer length factor
    answer_length_factor = min(len(answer) / 800, 1.0)
    
    # Specificity factor
    specificity_keywords = ["specifically", "for example", "including", "such as", "details", "code", "example"]
    specificity_factor = 1.0 if any(keyword in answer.lower() for keyword in specificity_keywords) else 0.7
    
    confidence = (avg_retrieval_score * 0.6 + answer_length_factor * 0.2 + specificity_factor * 0.2)
    
    return min(confidence, 0.95)

# API Routes
@app.get("/")
async def root():
    return {
        "message": "Course Q&A API - MongoDB Semantic Search", 
        "status": "running",
        "version": "4.8.0",
        "features": [
            "MongoDB Atlas Semantic Search",
            "Google Gemini AI Integration", 
            "Enhanced Context Retrieval",
            "Better Answer Relevance",
            "Advanced Citation System",
            "Performance Metrics & Analytics"
        ]
    }

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    status = {
        "mongo": check_database_available(),
        "embedding_model": embedding_model is not None,
        "gemini": gemini_model is not None,
        "services_initialized": services_initialized
    }
    
    overall_status = "healthy" if all([status["mongo"], status["embedding_model"], status["gemini"]]) else "degraded"
    if not status["mongo"]:
        overall_status = "critical"
    
    return {
        "status": overall_status, 
        "services": status,
        "version": "4.8.0"
    }

@app.post("/api/v1/upload", response_model=UploadResponse)
async def upload_file(file: UploadFile = File(...)):
    """Upload and process course material file"""
    upload_start_time = time.time()
    
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")
    
    if not check_database_available():
        raise HTTPException(status_code=503, detail="Database service temporarily unavailable")
    
    # Check file size (limit to 50MB)
    max_size = 50 * 1024 * 1024
    file.file.seek(0, 2)  # Seek to end
    file_size = file.file.tell()
    file.file.seek(0)  # Reset to beginning
    
    if file_size > max_size:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 50MB.")
    
    file_content = await file.read()
    
    if len(file_content) == 0:
        raise HTTPException(status_code=400, detail="Empty file provided")
    
    filename = file.filename.lower()
    
    try:
        # Extract text based on file type
        if filename.endswith('.pdf'):
            text = extract_text_from_pdf(file_content)
            processing_type = "PDF extraction"
        elif filename.endswith('.docx'):
            text = extract_text_from_docx(file_content)
            processing_type = "DOCX extraction"
        elif filename.endswith(('.xlsx', '.xls')):
            text = extract_text_from_xlsx(file_content)
            processing_type = "Excel extraction"
        elif filename.endswith('.csv'):
            text = extract_text_from_csv(file_content)
            processing_type = "CSV extraction"
        elif filename.endswith('.txt') or filename.endswith('.md'):
            text = extract_text_from_txt(file_content)
            processing_type = "Text extraction"
        else:
            raise HTTPException(status_code=400, detail="Unsupported file format. Supported formats: PDF, DOCX, XLSX, CSV, TXT, MD")
        
        if not text or len(text.strip()) == 0:
            raise HTTPException(status_code=400, detail="No text content could be extracted from the file")
        
        logger.info(f"📄 Extracted {len(text)} characters from {file.filename}")
        
        # Use improved chunking
        chunks = intelligent_chunking(text, file.filename, chunk_size=800, overlap=100)
        logger.info(f"📦 Created {len(chunks)} intelligent chunks from {file.filename}")
        
        if not chunks:
            raise HTTPException(status_code=400, detail="No meaningful content chunks could be created from the extracted text")
        
        # Process chunks
        processed_count = 0
        failed_count = 0
        
        for i, chunk in enumerate(chunks):
            try:
                # Generate embedding
                embedding = generate_embedding(chunk["content"])
                
                # Add embedding to chunk for semantic search
                chunk["embedding"] = embedding
                
                # Store in MongoDB
                documents_collection.insert_one(chunk)
                processed_count += 1
                
                logger.info(f"✅ Successfully stored chunk {chunk['chunk_id']} in MongoDB")
                
            except Exception as e:
                logger.error(f"Error processing chunk {i}: {e}")
                failed_count += 1
                continue
        
        processing_time = time.time() - upload_start_time
        PerformanceMetrics.track_upload(
            filename=file.filename,
            chunks_processed=processed_count,
            chunks_failed=failed_count,
            processing_time=processing_time
        )
        
        success_message = f"Successfully processed {file.filename}. Created {processed_count} chunks with semantic embeddings."
        if failed_count > 0:
            success_message += f" {failed_count} chunks failed to process."
        
        return UploadResponse(
            message=success_message,
            document_id=str(uuid.uuid4()),
            processing_type=processing_type,
            chunks_processed=processed_count,
            chunks_failed=failed_count
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error in file upload: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing file: {str(e)}")

@app.get("/api/v1/documents/count", response_model=DocumentCountResponse)
async def get_document_count():
    """Get count of processed documents"""
    try:
        if not check_database_available():
            return DocumentCountResponse(document_count=0)
        count = documents_collection.count_documents({})
        return DocumentCountResponse(document_count=count)
    except Exception as e:
        logger.error(f"Error getting document count: {e}")
        return DocumentCountResponse(document_count=0)

@app.delete("/api/v1/documents")
async def clear_all_documents():
    """Clear all documents from MongoDB"""
    try:
        if not check_database_available():
            return {"message": "Database service not available"}
        
        # Delete from MongoDB
        mongo_result = documents_collection.delete_many({})
        
        return {
            "message": f"Successfully cleared {mongo_result.deleted_count} documents from MongoDB"
        }
    except Exception as e:
        logger.error(f"Error clearing documents: {e}")
        raise HTTPException(status_code=500, detail=f"Error clearing documents: {str(e)}")

@app.get("/api/v1/answer", response_model=AnswerResponse)
async def get_answer(
    query: str,
    lang: str = "en",
    top_k: int = 15
):
    """Get comprehensive answer to course-related question"""
    start_time = time.time()
    
    try:
        if not check_database_available():
            raise HTTPException(status_code=503, detail="Database service temporarily unavailable")
        
        doc_count = documents_collection.count_documents({})
        
        if doc_count == 0:
            raise HTTPException(status_code=404, detail="No course materials found. Please upload documents first.")
        
        logger.info(f"🔍 Processing query: '{query}'")
        
        # Analyze question type with enhanced processing
        question_type = analyze_question_type(query)
        logger.info(f"📝 Question type: {question_type['type']}")
        
        # Perform enhanced question-aware retrieval using semantic search
        retrieved_chunks = question_aware_retrieval(query, question_type, top_k)
        logger.info(f"📚 Retrieved {len(retrieved_chunks)} relevant chunks")
        
        # Generate answer using enhanced Gemini with context
        answer_data = await generate_comprehensive_answer(query, retrieved_chunks, question_type)
        
        latency_ms = (time.time() - start_time) * 1000
        
        # Track performance metrics
        PerformanceMetrics.track_query(
            query=query,
            response_time=latency_ms,
            confidence=answer_data["confidence"],
            relevant_docs=len(retrieved_chunks),
            total_docs=doc_count,
            success=answer_data["confidence"] > 0.3 and len(retrieved_chunks) > 0
        )
        
        logger.info(f"✅ Answer generated in {latency_ms:.2f}ms with confidence {answer_data['confidence']:.2f}")
        
        return AnswerResponse(
            answer=answer_data["answer"],
            citations=answer_data["citations"],
            confidence=answer_data["confidence"],
            document_count=doc_count,
            relevant_docs=len(retrieved_chunks),
            latency_ms=latency_ms
        )
        
    except HTTPException as e:
        latency_ms = (time.time() - start_time) * 1000
        PerformanceMetrics.track_query(
            query=query,
            response_time=latency_ms,
            confidence=0.0,
            relevant_docs=0,
            total_docs=documents_collection.count_documents({}) if documents_collection else 0,
            success=False
        )
        raise e
    except Exception as e:
        latency_ms = (time.time() - start_time) * 1000
        PerformanceMetrics.track_query(
            query=query,
            response_time=latency_ms,
            confidence=0.0,
            relevant_docs=0,
            total_docs=documents_collection.count_documents({}) if documents_collection else 0,
            success=False
        )
        logger.error(f"Error generating answer: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating answer: {str(e)}")

@app.get("/api/v1/source/{source_id}", response_model=SourceResponse)
async def get_source(source_id: str):
    """Get source document by source_id"""
    try:
        if not check_database_available():
            raise HTTPException(status_code=503, detail="Database service temporarily unavailable")
            
        document = documents_collection.find_one({"chunk_id": source_id})
        if not document:
            raise HTTPException(status_code=404, detail="Source not found")
        
        return SourceResponse(
            text=document["content"],
            metadata=document.get("metadata", {}),
            source_id=source_id
        )
    except Exception as e:
        logger.error(f"Error retrieving source: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving source: {str(e)}")

# Metrics Endpoints
@app.get("/api/v1/metrics/kpis", response_model=KPIsResponse)
async def get_kpis(timeframe: int = 24):
    """Get Key Performance Indicators"""
    kpis = PerformanceMetrics.calculate_kpis(timeframe)
    kpis["timeframe_hours"] = timeframe
    return KPIsResponse(**kpis)

@app.get("/api/v1/metrics/system", response_model=SystemMetricsResponse)
async def get_system_metrics():
    """Get current system performance metrics"""
    metrics = PerformanceMetrics.get_system_metrics()
    return SystemMetricsResponse(**metrics)

@app.get("/api/v1/metrics/cost", response_model=CostMetricsResponse)
async def get_cost_metrics():
    """Get cost estimation metrics"""
    total_cost = CostCalculator.calculate_total_cost()
    total_queries = len(performance_metrics["query_history"])
    
    return CostMetricsResponse(
        total_estimated_cost_usd=total_cost,
        cost_per_query=total_cost / total_queries if total_queries > 0 else 0,
        total_queries=total_queries,
        cost_breakdown={
            "gemini_estimation": CostCalculator.GEMINI_COST_PER_1K_TOKENS,
            "embedding_estimation": CostCalculator.EMBEDDING_COST_PER_1K
        }
    )

@app.get("/api/v1/metrics/export")
async def export_metrics():
    """Export all metrics as CSV"""
    try:
        # Convert to pandas DataFrame for easy CSV export
        queries_df = pd.DataFrame(performance_metrics["query_history"]) if performance_metrics["query_history"] else pd.DataFrame()
        uploads_df = pd.DataFrame(performance_metrics["upload_history"]) if performance_metrics["upload_history"] else pd.DataFrame()
        
        # Create CSV responses
        queries_csv = queries_df.to_csv(index=False) if not queries_df.empty else "No query data available"
        uploads_csv = uploads_df.to_csv(index=False) if not uploads_df.empty else "No upload data available"
        
        return {
            "queries": queries_csv,
            "uploads": uploads_csv,
            "export_time": datetime.utcnow().isoformat()
        }
    except Exception as e:
        logger.error(f"Error exporting metrics: {e}")
        raise HTTPException(status_code=500, detail=f"Error exporting metrics: {str(e)}")

# Error handlers
@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    return JSONResponse(
        status_code=exc.status_code,
        content={"detail": exc.detail}
    )

@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    logger.error(f"Unhandled exception: {exc}")
    return JSONResponse(
        status_code=500,
        content={"detail": "Internal server error"}
    )

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=PORT)

sys.path.append(os.path.dirname(os.path.abspath(".\generate_metrics_report.py")))

def generate_metrics_report():
    """Generate comprehensive metrics report using actual backend data"""
    
    try:
        # Import the performance metrics from your main.py
        from main import performance_metrics, CostCalculator, PerformanceMetrics
        
        # Get actual metrics data
        query_history = performance_metrics.get("query_history", [])
        upload_history = performance_metrics.get("upload_history", [])
        
        # Calculate time-based metrics
        kpis = PerformanceMetrics.calculate_kpis(24)  # Last 24 hours
        system_metrics = PerformanceMetrics.get_system_metrics()
        total_cost = CostCalculator.calculate_total_cost()
        
        # Prepare report data from actual metrics
        report_data = {
            "summary": {
                "total_queries": len(query_history),
                "time_period": "24 hours",
                "success_rate": kpis.get("success_rate", 0),
                "avg_response_time": kpis.get("avg_response_time", 0),
                "total_cost_estimate": total_cost
            },
            "performance_metrics": {
                "response_time_distribution": [q.get("response_time_ms", 0) for q in query_history[-50:]],  # Last 50 queries
                "confidence_scores": [q.get("confidence", 0) for q in query_history[-50:] if q.get("confidence", 0) > 0],
                "queries_per_hour": calculate_queries_per_hour(query_history),
                "success_rates": [1 if q.get("success", False) else 0 for q in query_history[-20:]]
            },
            "cost_analysis": {
                "cost_per_query": total_cost / len(query_history) if query_history else 0,
                "main_components": ["Gemini API", "Embedding Generation"],
                "cost_breakdown": [
                    total_cost * 0.8,  # Estimate 80% for Gemini
                    total_cost * 0.2   # Estimate 20% for embeddings
                ]
            },
            "system_health": {
                "memory_usage_mb": system_metrics.get("memory_usage_mb", 0),
                "cpu_percent": system_metrics.get("cpu_percent", 0),
                "uptime_seconds": system_metrics.get("uptime_seconds", 0),
                "total_uploads": len(upload_history)
            }
        }
        
    except ImportError:
        print("⚠️  Could not import from main.py, using sample data")
        # Fallback to sample data if main.py is not available
        report_data = create_sample_data()
    
    # Create visualizations
    plt.style.use('default')
    fig, axes = plt.subplots(2, 2, figsize=(15, 12))
    
    # Response time distribution
    if report_data["performance_metrics"]["response_time_distribution"]:
        axes[0, 0].hist(report_data["performance_metrics"]["response_time_distribution"], bins=10, alpha=0.7, color='skyblue', edgecolor='black')
        axes[0, 0].set_title('Response Time Distribution (ms)')
        axes[0, 0].set_xlabel('Response Time (ms)')
        axes[0, 0].set_ylabel('Frequency')
        axes[0, 0].grid(True, alpha=0.3)
    else:
        axes[0, 0].text(0.5, 0.5, 'No response time data\navailable', ha='center', va='center', transform=axes[0, 0].transAxes)
        axes[0, 0].set_title('Response Time Distribution (ms)')
    
    # Confidence scores
    if report_data["performance_metrics"]["confidence_scores"]:
        axes[0, 1].hist(report_data["performance_metrics"]["confidence_scores"], bins=10, alpha=0.7, color='lightgreen', edgecolor='black')
        axes[0, 1].set_title('Confidence Score Distribution')
        axes[0, 1].set_xlabel('Confidence Score')
        axes[0, 1].set_ylabel('Frequency')
        axes[0, 1].grid(True, alpha=0.3)
    else:
        axes[0, 1].text(0.5, 0.5, 'No confidence score data\navailable', ha='center', va='center', transform=axes[0, 1].transAxes)
        axes[0, 1].set_title('Confidence Score Distribution')
    
    # Query volume over time
    if report_data["performance_metrics"]["queries_per_hour"]:
        hours = list(range(len(report_data["performance_metrics"]["queries_per_hour"])))
        axes[1, 0].plot(hours, report_data["performance_metrics"]["queries_per_hour"], marker='o', color='orange', linewidth=2, markersize=6)
        axes[1, 0].set_title('Query Volume Over Time')
        axes[1, 0].set_xlabel('Hour')
        axes[1, 0].set_ylabel('Queries per Hour')
        axes[1, 0].grid(True, alpha=0.3)
        axes[1, 0].fill_between(hours, report_data["performance_metrics"]["queries_per_hour"], alpha=0.3, color='orange')
    else:
        axes[1, 0].text(0.5, 0.5, 'No query volume data\navailable', ha='center', va='center', transform=axes[1, 0].transAxes)
        axes[1, 0].set_title('Query Volume Over Time')
    
    # Cost breakdown
    components = report_data["cost_analysis"]["main_components"]
    costs = report_data["cost_analysis"]["cost_breakdown"]
    if any(costs):
        axes[1, 1].pie(costs, labels=components, autopct='%1.1f%%', startangle=90, colors=['#ff9999','#66b3ff'])
        axes[1, 1].set_title('Cost Breakdown')
    else:
        axes[1, 1].text(0.5, 0.5, 'No cost data\navailable', ha='center', va='center', transform=axes[1, 1].transAxes)
        axes[1, 1].set_title('Cost Breakdown')
    
    plt.tight_layout()
    plt.savefig('metrics_report.png', dpi=300, bbox_inches='tight')
    plt.close()
    
    # Generate CSV report with actual query data
    try:
        if query_history:
            csv_data = {
                'timestamp': [q.get('timestamp', '') for q in query_history[-100:]],  # Last 100 queries
                'query': [q.get('query', '')[:100] + '...' if len(q.get('query', '')) > 100 else q.get('query', '') for q in query_history[-100:]],
                'response_time_ms': [q.get('response_time_ms', 0) for q in query_history[-100:]],
                'confidence': [q.get('confidence', 0) for q in query_history[-100:]],
                'relevant_docs': [q.get('relevant_docs', 0) for q in query_history[-100:]],
                'success': [q.get('success', False) for q in query_history[-100:]]
            }
            df = pd.DataFrame(csv_data)
            df.to_csv('performance_metrics.csv', index=False)
        else:
            # Create empty CSV with headers if no data
            df = pd.DataFrame(columns=['timestamp', 'query', 'response_time_ms', 'confidence', 'relevant_docs', 'success'])
            df.to_csv('performance_metrics.csv', index=False)
    except Exception as e:
        print(f"⚠️  Error generating CSV: {e}")
    
    # Generate summary report
    success_rate = report_data['summary']['success_rate']
    avg_response_time = report_data['summary']['avg_response_time']
    total_queries = report_data['summary']['total_queries']
    
    summary_report = f"""
    Course Q&A Chatbot - Performance Metrics Report
    Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
    
    SUMMARY:
    ---------
    Total Queries Processed: {total_queries}
    Time Period: {report_data['summary']['time_period']}
    Success Rate: {success_rate:.1%}
    Average Response Time: {avg_response_time:.0f} ms
    Total Estimated Cost: ${report_data['summary']['total_cost_estimate']:.4f}
    
    SYSTEM HEALTH:
    -------------
    Memory Usage: {report_data['system_health']['memory_usage_mb']:.1f} MB
    CPU Usage: {report_data['system_health']['cpu_percent']:.1f}%
    Uptime: {report_data['system_health']['uptime_seconds'] / 3600:.1f} hours
    Total Documents Uploaded: {report_data['system_health']['total_uploads']}
    
    PERFORMANCE INDICATORS:
    ----------------------
    - Average confidence score: {np.mean(report_data['performance_metrics']['confidence_scores']) if report_data['performance_metrics']['confidence_scores'] else 0:.1%}
    - Peak query volume: {max(report_data['performance_metrics']['queries_per_hour']) if report_data['performance_metrics']['queries_per_hour'] else 0} queries/hour
    - Total successful queries: {int(success_rate * total_queries) if total_queries > 0 else 0}
    
    COST ANALYSIS:
    -------------
    Cost per Query: ${report_data['cost_analysis']['cost_per_query']:.4f}
    Most expensive component: {components[costs.index(max(costs))] if costs else 'N/A'}
    
    RECOMMENDATIONS:
    ---------------
    1. {'Optimize query processing for faster responses' if avg_response_time > 3000 else 'Response times are optimal'}
    2. {'Improve retrieval accuracy' if success_rate < 0.8 else 'Success rate is good'}
    3. {'Monitor cost efficiency' if report_data['cost_analysis']['cost_per_query'] > 0.01 else 'Cost efficiency is good'}
    """
    
    with open('metrics_summary.txt', 'w') as f:
        f.write(summary_report)
    
    print("✅ Metrics report generated successfully!")
    print(f"   - Total queries analyzed: {total_queries}")
    print(f"   - Success rate: {success_rate:.1%}")
    print(f"   - Average response time: {avg_response_time:.0f} ms")
    print("   - metrics_report.png (Visualizations)")
    print("   - performance_metrics.csv (Raw data)")
    print("   - metrics_summary.txt (Summary report)")

def calculate_queries_per_hour(query_history):
    """Calculate queries per hour from query history"""
    if not query_history:
        return []
    
    # Group queries by hour
    hourly_counts = {}
    for query in query_history[-24:]:  # Last 24 queries or adjust as needed
        try:
            timestamp = datetime.fromisoformat(query.get('timestamp', ''))
            hour_key = timestamp.strftime('%H:00')
            hourly_counts[hour_key] = hourly_counts.get(hour_key, 0) + 1
        except:
            continue
    
    # Return counts for visualization
    return list(hourly_counts.values())[-8:]  # Last 8 hours

def create_sample_data():
    """Create sample data for demonstration"""
    return {
        "summary": {
            "total_queries": 45,
            "time_period": "24 hours",
            "success_rate": 0.82,
            "avg_response_time": 1850,
            "total_cost_estimate": 0.023
        },
        "performance_metrics": {
            "response_time_distribution": [1200, 1800, 2200, 1600, 2500, 1900, 2100],
            "confidence_scores": [0.7, 0.8, 0.9, 0.6, 0.85, 0.75, 0.88],
            "queries_per_hour": [5, 8, 12, 7, 10, 15, 9, 11],
            "success_rates": [1, 1, 1, 0, 1, 1, 1, 1, 0, 1]
        },
        "cost_analysis": {
            "cost_per_query": 0.0005,
            "main_components": ["Gemini API", "Embedding Generation"],
            "cost_breakdown": [0.0004, 0.0001]
        },
        "system_health": {
            "memory_usage_mb": 245.6,
            "cpu_percent": 12.3,
            "uptime_seconds": 86400,
            "total_uploads": 8
        }
    }

if __name__ == "__main__":
    generate_metrics_report()